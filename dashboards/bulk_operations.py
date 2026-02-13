"""
Bulk import/export operations for students, exams, questions
"""
import csv
import io
from django.utils import timezone
from django.db import transaction
from accounts.models import User
from exams.models import Exam, Question, Choice, SchoolClass, Subject
from .models import (
    BulkImportJob, BulkExportJob, QuestionBank, 
    QuestionCategory, QuestionTag, StudentPerformance
)


class BulkImporter:
    """Handle bulk imports from CSV"""
    
    @staticmethod
    def import_students(csv_file, school_class, created_by):
        """Import students from CSV"""
        job = BulkImportJob.objects.create(
            import_type='students',
            csv_file=csv_file,
            created_by=created_by,
            status='processing',
            started_at=timezone.now(),
        )
        
        errors = []
        successes = 0
        
        try:
            file_content = csv_file.read().decode('utf-8')
            reader = csv.DictReader(io.StringIO(file_content))
            
            required_fields = ['first_name', 'last_name', 'username', 'email']
            if reader.fieldnames and not all(f in reader.fieldnames for f in required_fields):
                job.status = 'failed'
                job.error_log = f"Missing required fields. Expected: {', '.join(required_fields)}"
                job.completed_at = timezone.now()
                job.save()
                return job
            
            for row_num, row in enumerate(reader, start=2):
                try:
                    username = row.get('username', '').strip()
                    email = row.get('email', '').strip()
                    first_name = row.get('first_name', '').strip()
                    last_name = row.get('last_name', '').strip()
                    
                    if not all([username, email, first_name, last_name]):
                        errors.append(f"Row {row_num}: Missing required field")
                        continue
                    
                    if User.objects.filter(username=username).exists():
                        errors.append(f"Row {row_num}: Username '{username}' already exists")
                        continue
                    
                    if User.objects.filter(email=email).exists():
                        errors.append(f"Row {row_num}: Email '{email}' already exists")
                        continue
                    
                    user = User.objects.create_user(
                        username=username,
                        email=email,
                        first_name=first_name,
                        last_name=last_name,
                        role=User.Role.STUDENT,
                        student_class=school_class,
                    )
                    successes += 1
                
                except Exception as e:
                    errors.append(f"Row {row_num}: {str(e)}")
            
            job.status = 'completed'
            job.total_rows = successes + len(errors)
            job.successful_rows = successes
            job.failed_rows = len(errors)
            job.error_log = '\n'.join(errors) if errors else ''
            job.completed_at = timezone.now()
            job.save()
            
        except Exception as e:
            job.status = 'failed'
            job.error_log = str(e)
            job.completed_at = timezone.now()
            job.save()
        
        return job
    
    @staticmethod
    def import_questions(csv_file, created_by, category_id=None):
        """Import questions from CSV, Excel, or Word"""
        job = BulkImportJob.objects.create(
            import_type='questions',
            csv_file=csv_file,
            created_by=created_by,
            status='processing',
            started_at=timezone.now(),
        )
        
        errors = []
        successes = 0
        
        try:
            if hasattr(csv_file, 'seek'):
                csv_file.seek(0)

            filename = csv_file.name.lower()
            questions_data = []
            
            # Determine parser based on file extension
            if filename.endswith('.xlsx'):
                questions_data, parse_errors = BulkImporter._parse_excel_questions(csv_file)
                errors.extend(parse_errors)
            elif filename.endswith('.docx'):
                questions_data, parse_errors = BulkImporter._parse_word_questions(csv_file)
                errors.extend(parse_errors)
            else:
                # Default to CSV
                questions_data, parse_errors = BulkImporter._parse_csv_questions(csv_file)
                errors.extend(parse_errors)
            
            if not questions_data and not errors:
                 errors.append("No data found in file")

            category = None
            if category_id:
                category = QuestionCategory.objects.get(id=category_id)
            
            for row_num, row in enumerate(questions_data, start=1):
                try:
                    text = row.get('text', '').strip()
                    q_type = row.get('type', '').lower()
                    marks = int(row.get('marks', '1'))
                    class_name = row.get('class', '').strip()
                    subject_name = row.get('subject', '').strip()

                    if not text or q_type not in ['objective', 'subjective', 'short_answer']:
                        errors.append(f"Row {row_num}: Invalid question type or missing text")
                        continue

                    # Lookup Class and Subject
                    school_class = None
                    if class_name:
                        try:
                            school_class = SchoolClass.objects.get(name__iexact=class_name)
                        except SchoolClass.DoesNotExist:
                            errors.append(f"Row {row_num}: Class '{class_name}' not found")
                            continue

                    subject = None
                    if subject_name:
                        try:
                            subject = Subject.objects.get(name__iexact=subject_name)
                        except Subject.DoesNotExist:
                            errors.append(f"Row {row_num}: Subject '{subject_name}' not found")
                            continue
                    
                    question = QuestionBank.objects.create(
                        text=text,
                        question_type=q_type,
                        marks=marks,
                        difficulty='medium',  # Default since removed from template
                        category=category,
                        school_class=school_class,
                        subject=subject,
                        created_by=created_by,
                        is_published=True,
                    )
                    
                    # Handle choices for objective questions
                    if q_type == 'objective':
                        for i in range(1, 5):
                            choice_text = row.get(f'choice_{i}', '')
                            is_correct_val = str(row.get(f'correct_{i}', '')).lower().strip()
                            is_correct = is_correct_val in ['true', '1', 'yes']
                            
                            if choice_text:
                                from .models import QuestionChoice
                                QuestionChoice.objects.create(
                                    question=question,
                                    text=choice_text,
                                    is_correct=is_correct,
                                    order=i,
                                )
                    
                    successes += 1
                
                except Exception as e:
                    errors.append(f"Row {row_num}: {str(e)}")
            
            job.status = 'completed' if not (successes == 0 and len(errors) > 0) else 'failed'
            job.total_rows = successes + len(errors)
            job.successful_rows = successes
            job.failed_rows = len(errors)
            job.error_log = '\n'.join(errors) if errors else ''
            job.completed_at = timezone.now()
            job.save()
            
        except Exception as e:
            job.status = 'failed'
            job.error_log = str(e)
            job.completed_at = timezone.now()
            job.save()
        
        return job

    @staticmethod
    def _parse_csv_questions(file_obj):
        """Parse CSV content"""
        data = []
        errors = []
        try:
            file_content = file_obj.read().decode('utf-8')
            reader = csv.DictReader(io.StringIO(file_content))
            
            required = ['text', 'type', 'marks', 'class', 'subject']
            if reader.fieldnames and not all(f in reader.fieldnames for f in required):
                return [], [f"CSV Missing required fields: {', '.join(required)}"]

            for row in reader:
                data.append(row)
        except Exception as e:
            errors.append(f"CSV Parse Error: {str(e)}")
        return data, errors

    @staticmethod
    def _parse_excel_questions(file_obj):
        """Parse Excel content using openpyxl"""
        import openpyxl
        data = []
        errors = []
        try:
            wb = openpyxl.load_workbook(file_obj, data_only=True)
            sheet = wb.active
            
            # Read header
            headers = [cell.value for cell in sheet[1]]
            required = ['text', 'type', 'marks', 'class', 'subject']
            if not all(field in headers for field in required):
                 return [], [f"Excel Missing required headers: {', '.join(required)}"]

            # Map headers to indices
            header_map = {h: i for i, h in enumerate(headers) if h}

            for row_idx, row in enumerate(sheet.iter_rows(min_row=2), start=2):
                row_data = {}
                has_content = False
                for field, col_idx in header_map.items():
                    val = row[col_idx].value
                    if val is not None:
                         row_data[field] = str(val).strip()
                         has_content = True
                    else:
                         row_data[field] = ""
                
                if has_content:
                    data.append(row_data)
        except Exception as e:
            errors.append(f"Excel Parse Error: {str(e)}")
        return data, errors

    @staticmethod
    def _parse_word_questions(file_obj):
        """Parse Word content using python-docx (Table based)"""
        import docx
        data = []
        errors = []
        try:
            doc = docx.Document(file_obj)
            if not doc.tables:
                return [], ["No tables found in Word document"]
            
            table = doc.tables[0]
            if not table.rows:
                 return [], ["Empty table in Word document"]

            # Header
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            required = ['text', 'type', 'marks', 'class', 'subject']
            if not all(field in headers for field in required):
                 return [], [f"Word Table Missing required headers: {', '.join(required)}"]
            
            header_map = {h: i for i, h in enumerate(headers) if h}

            for row_idx, row in enumerate(table.rows[1:], start=2):
                row_data = {}
                # Ensure row has enough cells
                has_content = False
                for field, col_idx in header_map.items():
                    if col_idx < len(row.cells):
                        val = row.cells[col_idx].text.strip()
                        row_data[field] = val
                        if val: has_content = True
                    else:
                        row_data[field] = ""
                
                if has_content:
                    data.append(row_data)

        except Exception as e:
            errors.append(f"Word Parse Error: {str(e)}")
        return data, errors


class BulkExporter:
    """Handle bulk exports to CSV/Excel"""
    
    @staticmethod
    def export_exam_results(exam, format_type='csv'):
        """Export exam results to CSV/Excel"""
        from exams.models import ExamAttempt
        
        job = BulkExportJob.objects.create(
            export_type='results',
            exam=exam,
            exported_by=None,  # Set by view
            status='processing',
            file_format=format_type,
        )
        
        try:
            attempts = ExamAttempt.objects.filter(
                exam=exam,
                status='submitted'
            ).select_related('student')
            
            # Create CSV in memory
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Header
            writer.writerow([
                'Student ID', 'Student Name', 'Email',
                'Score', 'Total Marks', 'Percentage',
                'Time Taken (mins)', 'Status', 'Submitted At'
            ])
            
            # Data rows
            for attempt in attempts:
                time_minutes = (attempt.completed_at - attempt.started_at).total_seconds() / 60 if attempt.completed_at else 0
                total_marks = sum([q.marks for q in exam.questions.all()])
                percentage = (attempt.total_score / total_marks * 100) if total_marks > 0 else 0
                
                writer.writerow([
                    attempt.student.id,
                    f"{attempt.student.first_name} {attempt.student.last_name}",
                    attempt.student.email,
                    attempt.total_score,
                    total_marks,
                    f"{percentage:.2f}",
                    f"{time_minutes:.2f}",
                    'Passed' if percentage >= 60 else 'Failed',
                    attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else '',
                ])
            
            csv_content = output.getvalue()
            
            # Save file
            from django.core.files.base import ContentFile
            filename = f"results_{exam.id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.csv"
            job.export_file.save(
                filename,
                ContentFile(csv_content.encode('utf-8')),
                save=True
            )
            
            job.status = 'completed'
            job.completed_at = timezone.now()
            job.save()
            
        except Exception as e:
            job.status = 'failed'
            job.completed_at = timezone.now()
            job.save()
        
        return job
    
    @staticmethod
    def export_student_performance(student, format_type='csv'):
        """Export student performance report"""
        job = BulkExportJob.objects.create(
            export_type='performance',
            exported_by=None,  # Set by view
            status='processing',
            file_format=format_type,
        )
        
        try:
            performances = StudentPerformance.objects.filter(
                student=student
            ).select_related('exam').order_by('-attempted_at')
            
            output = io.StringIO()
            writer = csv.writer(output)
            
            writer.writerow([
                'Exam', 'Attempt', 'Score', 'Total Marks',
                'Percentage', 'Time Taken (mins)', 'Status', 'Date'
            ])
            
            for perf in performances:
                time_minutes = perf.time_taken / 60
                writer.writerow([
                    perf.exam.title,
                    perf.attempt_number,
                    perf.score,
                    perf.total_marks,
                    f"{perf.percentage:.2f}",
                    f"{time_minutes:.2f}",
                    perf.status.capitalize(),
                    perf.attempted_at.strftime('%Y-%m-%d %H:%M:%S'),
                ])
            
            csv_content = output.getvalue()
            
            filename = f"performance_{student.id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.csv"
            from django.core.files.base import ContentFile
            job.export_file.save(
                filename,
                ContentFile(csv_content.encode('utf-8')),
                save=True
            )
            
            job.status = 'completed'
            job.completed_at = timezone.now()
            job.save()
            
        except Exception as e:
            job.status = 'failed'
            job.completed_at = timezone.now()
            job.save()
        
        return job
